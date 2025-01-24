import streamlit as st
import docx
from docx.shared import Pt
import tempfile
import os

def generate_document(details):
    doc = docx.Document()

    doc.add_heading('RENTAL AGREEMENT', 0)

    date_of_agreement = f"{details['What is the date of the agreement (DD)?']}/{details['What is the month of the agreement (MM)?']}/{details['What is the year of the agreement (YYYY)?']}"
    
    doc.add_paragraph(
        f"This RENTAL AGREEMENT is executed at {details['What is the place of the agreement?']} "
        f"on this {details['What is the date of the agreement (DD)?']} day of {details['What is the month of the agreement (MM)?']} "
        f"{details['What is the year of the agreement (YYYY)?']}."
    )

    owner_name = details["What is the owner's name?"]
    owner_father_name = details["What is the owner's father's name?"]
    owner_address = details["What is the owner's permanent address?"]

    tenant_name = details["What is the tenant's name?"]
    tenant_father_name = details["What is the tenant's father's name?"]
    tenant_address = details["What is the tenant's permanent address?"]

    doc.add_paragraph(
        f"BETWEEN {owner_name}, S/O {owner_father_name}, "
        f"residing at {owner_address} (hereinafter referred to as 'OWNER') "
        f"AND {tenant_name}, S/O {tenant_father_name}, "
        f"residing at {tenant_address} (hereinafter referred to as 'TENANT')."
    )

    doc.add_paragraph(
        "WHEREAS, the OWNER has agreed to rent out the property described below to the TENANT, "
        "and the TENANT has agreed to take the same on rent, on the terms and conditions set forth herein."
    )

    doc.add_heading('I. PREMISES', level=1)
    doc.add_paragraph(
        f"The property being rented is located at {details['What is the complete address of the rented property?']} and consists of "
        f"{details['What is the number of bedrooms?']} bedrooms, {details['What is the number of fans?']} fans, "
        f"{details['What is the number of CFL lights?']} CFL lights, {details['What is the number of geysers?']} geysers, "
        f"and {details['What is the number of mirrors?']} mirrors."
    )

    doc.add_heading('II. TERM', level=1)
    doc.add_paragraph(
        f"The term of this Agreement shall commence on {details['What is the starting date of the agreement (DD/MM/YYYY)?']} "
        f"and shall continue until {details['What is the expiry date of the agreement (DD/MM/YYYY)?']}."
    )

    doc.add_heading('III. RENT', level=1)
    doc.add_paragraph(
        f"The TENANT shall pay a monthly rent of Rs. {details['What is the monthly rent amount?']} to the OWNER."
    )

    doc.add_heading('IV. MAINTENANCE', level=1)
    doc.add_paragraph(
        f"The TENANT shall pay a monthly maintenance charge of Rs. {details['What is the monthly maintenance charge?']}."
    )

    doc.add_heading('V. SECURITY DEPOSIT', level=1)
    doc.add_paragraph(
        f"The TENANT has paid a security deposit of Rs. {details['What is the security deposit amount?']} "
        f"via cheque number {details['What is the cheque number?']} dated {details['What is the cheque date (DD/MM/YYYY)?']}."
    )

    doc.add_heading('VI. TERMINATION', level=1)
    doc.add_paragraph(
        f"This Agreement may be terminated by either party on {details['What is the number of months notice required for termination?']} months' written notice to the other party, "
        "upon expiry of the term or upon the happening of any of the events contemplated under Clause V above."
    )

    doc.add_heading('VII. INDEMNITY', level=1)
    doc.add_paragraph(
        "The TENANT shall indemnify and hold harmless the OWNER against all liabilities, claims, damages, costs and expenses whatsoever, "
        "arising out of the use of the rented premises by the TENANT or any guest of the TENANT, whether by contract, tort or otherwise."
    )

    doc.add_heading('VIII. NOTICES', level=1)
    doc.add_paragraph(
        "Any notice or communication under this Agreement shall be deemed to have been received by the intended party if it is sent by registered post, "
        "or by hand or messenger or by e-mail, and any such notice shall be deemed to have been served at the time of delivery."
    )

    doc.add_heading('IX. BREACH OF AGREEMENT', level=1)
    doc.add_paragraph(
        "If the TENANT breaches any of the terms and conditions of this Agreement, the OWNER may, after giving "
        "a reasonable opportunity of being heard, summarily expel the TENANT from the rented premises. In case the OWNER "
        "is not able to obtain the rented premises, the TENANT shall be responsible for paying compensation for any damage caused to "
        "the rented premises or to any property of the OWNER in connection with such eviction."
    )

    doc.add_heading('X. WITNESSES', level=1)
    doc.add_paragraph(
        f"We hereby appoint {details['What is the name of the first witness?']} of {details['What is the address of the first witness?']} and "
        f"{details['What is the name of the second witness?']} of {details['What is the address of the second witness?']} as witnesses to this RENTAL AGREEMENT. "
        "We both sign below our names."
    )

    doc.add_paragraph(f"Date: {date_of_agreement}")

    doc.add_paragraph(f"For OWNER: {owner_name}")
    doc.add_paragraph(f"For TENANT: {tenant_name}")

    # Save the document in a temporary directory
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, 'Rental_Agreement.docx')
    doc.save(file_path)
    return file_path

def main():
    st.title("Rental Agreement Form")
    st.write("Please fill in the details below to generate the rental agreement document.")

    questions = [
        "What is the place of the agreement?",
        "What is the date of the agreement (DD)?",
        "What is the month of the agreement (MM)?",
        "What is the year of the agreement (YYYY)?",
        "What is the owner's name?",
        "What is the owner's father's name?",
        "What is the owner's permanent address?",
        "What is the tenant's name?",
        "What is the tenant's father's name?",
        "What is the tenant's office/institution address?",
        "What is the tenant's permanent address?",
        "What is the complete address of the rented property?",
        "What is the starting date of the agreement (DD/MM/YYYY)?",
        "What is the expiry date of the agreement (DD/MM/YYYY)?",
        "What is the monthly rent amount?",
        "What is the monthly maintenance charge?",
        "What is the security deposit amount?",
        "What is the cheque number?",
        "What is the cheque date (DD/MM/YYYY)?",
        "What is the number of bedrooms?",
        "What is the number of fans?",
        "What is the number of CFL lights?",
        "What is the number of geysers?",
        "What is the number of mirrors?",
        "What is the name of the first witness?",
        "What is the address of the first witness?",
        "What is the name of the second witness?",
        "What is the address of the second witness?",
        "What is the number of months notice required for termination?"
    ]

    responses = {}
    for question in questions:
        responses[question] = st.text_input(question)

    if st.button("Generate Agreement"):
        file_path = generate_document(responses)
        st.success("Rental agreement created!")
        st.download_button(
            label="Download Rental Agreement",
            data=open(file_path, "rb"),
            file_name="Rental_Agreement.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()